import datetime
import threading
import asyncio
import pyaudio
import json
import tkinter as tk
from tkinter import ttk
from openai_chat import gen
from brainstorm_chat import brainstorm, end_brainstorm, is_brainstorming, get_brainstorm_summary, set_gui_callback, set_assistant_resume_callback
from assistantTools import get_events, get_free_slots_today, get_free_slots_week, get_todays_events, format_time, get_weeks_events, set_user_timezone, format_event_time
from deepgram import DeepgramClient, DeepgramClientOptions, LiveTranscriptionEvents, LiveOptions
from google_oauth import is_authenticated, authenticate, revoke_access
import os
from dotenv import load_dotenv
from calendar_events import create_event_from_speech, find_events_by_title, update_calendar_event, delete_calendar_event, parse_move_request, move_event_to_new_time
from pathlib import Path
import pytz
from embedded_config import EmbeddedConfig
from user_profile import UserProfile
import ssl
import certifi
import time

# Load environment variables from .env file
load_dotenv()

# Add this to your imports in assistant.py


# Global variable to store the assistant instance
current_assistant = None

# Initialize user profile for per-user custom commands
user_profile = UserProfile()

def load_custom_commands():
    """Load custom commands from user profile"""
    return user_profile.load_custom_commands()

def check_custom_commands(transcript):
    """Check if transcript matches any custom commands"""
    custom_commands = user_profile.load_custom_commands()
    transcript_lower = transcript.lower().strip()
    
    for base_command, custom_phrases in custom_commands.items():
        for phrase in custom_phrases:
            if phrase.lower().strip() in transcript_lower:
                print(f"🎯 Custom command matched: '{phrase}' -> '{base_command}'")
                return base_command
    
    return None

def display_brainstorm_message(message):
    """Display brainstorming summaries in the main GUI transcript"""
    try:
        # Access the global text_display widget
        import tkinter as tk
        
        # Find the text_display widget from globals or current assistant
        text_widget = None
        if 'text_display' in globals() and text_display:
            text_widget = text_display
        elif current_assistant and hasattr(current_assistant, 'text_display'):
            text_widget = current_assistant.text_display
        
        if text_widget:
            # Add timestamp for brainstorm messages
            import datetime
            timestamp = datetime.datetime.now().strftime("%H:%M:%S")
            
            # Add separator for brainstorm messages
            text_widget.insert(tk.END, f"\n🧠 [{timestamp}] " + "="*50 + "\n")
            text_widget.insert(tk.END, f"{message}\n")
            text_widget.insert(tk.END, "="*60 + "\n\n")
            text_widget.see(tk.END)
            print("📤 Brainstorm summary sent to GUI")
        else:
            print("⚠️ GUI text display not available")
    except Exception as e:
        print(f"❌ Error displaying brainstorm message in GUI: {e}")
        import traceback
        print(f"Full traceback: {traceback.format_exc()}")

class VoiceAssistant:
    def __init__(self, text_display):
        self.text_display = text_display
        self.deepgram = None
        self.dg_connection = None
        self.audio_stream = None
        self.is_running = False
        self.is_paused = False  # Add pause flag for brainstorming sessions
        self.loop = None
        self.waiting_for_question = False
        
        # Speaker diarization settings
        self.primary_speaker_id = None  # Track the primary user
        self.speaker_lock_enabled = True
        self.min_words_to_lock = 3  # Minimum words before locking speaker
        
        # Initialize Deepgram client with SSL context
        try:
            print("🔧 Initializing Deepgram client...")
            api_key = EmbeddedConfig.get_deepgram_key()
            if not api_key:
                print("❌ Warning: DEEPGRAM_API_KEY not found")
                return
            
            print(f"🔑 API Key found (length: {len(api_key)})") 
            
            # Use SSL context for certificate verification
            ssl_context = ssl.create_default_context(cafile=certifi.where())
            
            config = DeepgramClientOptions(
                options={
                    "keepalive": "true",
                    "ssl_context": ssl_context
                }
            )
            self.deepgram = DeepgramClient(api_key, config)
            print("✅ Deepgram client initialized successfully with SSL context")
        except Exception as e:
            print(f"❌ Error initializing Deepgram client: {e}")
        
    def filter_by_primary_speaker(self, result):
        """Filter transcript to only include primary speaker's words"""
        if not hasattr(result.channel.alternatives[0], 'words') or not result.channel.alternatives[0].words:
            print("🔍 No word-level data available, using full transcript")
            return result.channel.alternatives[0].transcript
        
        words = result.channel.alternatives[0].words
        
        # Group words by speaker
        speaker_words = {}
        for word_info in words:
            speaker_id = getattr(word_info, 'speaker', 0)
            if speaker_id not in speaker_words:
                speaker_words[speaker_id] = []
            speaker_words[speaker_id].append(word_info)
        
        print(f"🎤 Detected speakers: {list(speaker_words.keys())}")
        
        # Lock onto first speaker if not already locked
        if self.primary_speaker_id is None and self.speaker_lock_enabled:
            # Find speaker with most words in this utterance
            if speaker_words:
                primary_candidate = max(speaker_words.keys(), 
                                      key=lambda s: len(speaker_words[s]))
                
                # Only lock if they said enough words
                if len(speaker_words[primary_candidate]) >= self.min_words_to_lock:
                    self.primary_speaker_id = primary_candidate
                    print(f"🔒 Locked onto speaker {self.primary_speaker_id}")
                    self.text_display.insert(tk.END, f"🔒 Voice locked to Speaker {self.primary_speaker_id}\n")
                    self.text_display.see(tk.END)
                    # Update UI if function exists
                    try:
                        update_speaker_lock_status(True, self.primary_speaker_id)
                    except NameError:
                        pass
        
        # Return only primary speaker's words
        if self.primary_speaker_id is not None and self.primary_speaker_id in speaker_words:
            primary_words = speaker_words[self.primary_speaker_id]
            # Reconstruct transcript from primary speaker's words only
            filtered_transcript = ' '.join([getattr(word, 'word', '') for word in primary_words])
            print(f"👤 Primary speaker {self.primary_speaker_id} said: '{filtered_transcript}'")
            return filtered_transcript
        elif self.primary_speaker_id is not None:
            print(f"🔇 Filtered out - not primary speaker (locked to {self.primary_speaker_id})")
            return None
        
        # Fallback to full transcript if no speaker lock
        return result.channel.alternatives[0].transcript
    
    def reset_speaker_lock(self):
        """Reset speaker lock to re-identify primary speaker"""
        self.primary_speaker_id = None
        print("🔓 Speaker lock reset")
        self.text_display.insert(tk.END, "🔓 Speaker lock reset - will re-identify on next speech\n")
        self.text_display.see(tk.END)
        # Update UI if function exists
        try:
            update_speaker_lock_status(False)
        except NameError:
            pass
        
    def process_transcript(self, result):
        """Process the recognized text from Deepgram with speaker filtering"""
        try:
            # Skip processing if assistant is paused (during brainstorming)
            if self.is_paused:
                print("⏸️ Assistant paused - skipping transcript processing")
                return
            
            print("✅ Assistant active - processing transcript")
            
            # Extract text from Deepgram result
            if result.is_final:
                # Apply speaker filtering
                filtered_transcript = self.filter_by_primary_speaker(result)
                
                if filtered_transcript and filtered_transcript.strip():  # Only process non-empty transcripts
                    self.text_display.insert(tk.END, f"{filtered_transcript}\n")
                    self.text_display.see(tk.END)
                    print(f"Final processed transcript: {filtered_transcript}")
                    
                    # Check if we're waiting for a question response
                    if hasattr(self, 'waiting_for_question') and self.waiting_for_question:
                        self.handle_question_response(filtered_transcript)
                        return
                    
                    # First check for custom commands
                    custom_command = check_custom_commands(filtered_transcript)
                    if custom_command:
                        # Process the custom command as if it was the original command
                        self.text_display.insert(tk.END, f"🎯 Using custom command: {custom_command}\n")
                        filtered_transcript = custom_command  # Replace with base command
                    
                    # Handle different voice commands (including custom commands)
                    if "question" in filtered_transcript.lower():
                        self.handle_question_mode()
                    elif "brainstorm" in filtered_transcript.lower():
                        self.handle_brainstorm_session(filtered_transcript)
                    elif "when am i free this week" in filtered_transcript.lower():
                        self.handle_free_time_week()
                    elif "when am i free today" in filtered_transcript.lower():
                        self.handle_free_time_today()
                    elif "what events do i have today" in filtered_transcript.lower():
                        self.handle_todays_events()
                    elif "what events do i have this week" in filtered_transcript.lower():
                        self.handle_weeks_events()
                    elif "what events do i have coming up" in filtered_transcript.lower():
                        self.handle_upcoming_events()
                    elif "terminate" in filtered_transcript.lower():
                        self.stop_assistant()
                    elif "test resume" in filtered_transcript.lower():
                        # Test command to verify resume functionality
                        self.text_display.insert(tk.END, f"🧪 Testing pause/resume: is_paused = {self.is_paused}\n")
                        self.text_display.see(tk.END)
                        print(f"🧪 Testing pause/resume: is_paused = {self.is_paused}")
                    elif "create event" in filtered_transcript.lower() or "schedule" in filtered_transcript.lower():
                        self.handle_create_event(filtered_transcript)
                    elif "find event" in filtered_transcript.lower() or "search event" in filtered_transcript.lower():
                        search_term = filtered_transcript.lower().replace("find event", "").replace("search event", "").strip()
                        self.handle_find_event(search_term)
                    elif "delete event" in filtered_transcript.lower() or "cancel event" in filtered_transcript.lower():
                        search_term = filtered_transcript.lower().replace("delete event", "").replace("cancel event", "").strip()
                        self.handle_delete_event(search_term)
                    elif "move event" in filtered_transcript.lower() or "reschedule event" in filtered_transcript.lower():
                        self.handle_move_event(filtered_transcript)
                elif filtered_transcript is None:
                    # This means the speech was filtered out (not primary speaker)
                    pass  # Already logged in filter function
                    
        except Exception as e:
            print(f"Error processing transcript: {e}")
            import traceback
            print(f"Full traceback: {traceback.format_exc()}")
    
    def on_message(self, result, **kwargs):
        """Handle Deepgram message events"""
        try:
            sentence = result.channel.alternatives[0].transcript
            if len(sentence) == 0:
                return
            print(f"🔧 AUDIO DEBUG: Received transcript in main assistant: '{sentence}'")
            print(f"🔧 AUDIO DEBUG: is_paused = {self.is_paused}, is_final = {result.is_final}")
            self.process_transcript(result)
        except Exception as e:
            print(f"Error in on_message: {e}")
    
    def on_error(self, error, **kwargs):
        """Handle Deepgram error events"""
        print(f"Deepgram error: {error}")
    
    def on_close(self, close, **kwargs):
        """Handle Deepgram connection close"""
        print(f"Deepgram connection closed: {close}")
    
    def check_authentication(self):
        """Check if user is authenticated with Google Calendar"""
        if not is_authenticated():
            self.text_display.insert(tk.END, "❌ Google Calendar not authenticated. Please sign in first.\n")
            return False
        return True
    
    def handle_question_mode(self):
        """Handle question asking mode"""
        self.text_display.insert(tk.END, "Please ask your question...\n")
        print("Please ask your question...")
        
        # Set up a temporary callback to capture the next transcript
        self.waiting_for_question = True
    
    def handle_brainstorm_session(self, filtered_transcript):
        """Handle starting a brainstorming session - pauses current assistant instead of stopping"""
        self.text_display.insert(tk.END, "🧠 Starting brainstorming session - pausing assistant...\n")
        self.text_display.see(tk.END)
        print("🧠 Starting brainstorming session - pausing assistant...")
        
        # Pause the current assistant's audio processing instead of stopping completely
        self.is_paused = True  # Add a pause flag instead of stopping
        print("⏸️ Main assistant paused for brainstorming session")
        
        # Extract topic from the command (remove the trigger phrase)
        topic = filtered_transcript.lower().replace("brainstorm", "").replace("about", "").strip()
        
        # Start the voice brainstorming session
        try:
            response = brainstorm(topic if topic else None)
            self.text_display.insert(tk.END, f"{response}\n")
            self.text_display.see(tk.END)
            print(f"Brainstorm response: {response}")
        except Exception as e:
            error_msg = f"❌ Error starting brainstorming session: {str(e)}"
            self.text_display.insert(tk.END, f"{error_msg}\n")
            self.text_display.see(tk.END)
            print(error_msg)
            # If brainstorming fails, resume the main assistant
            self.is_paused = False
            print("▶️ Main assistant resumed due to brainstorming error")
    
    def resume_assistant(self):
        """Resume the main assistant after brainstorming session ends"""
        print(f"🔧 RESUME DEBUG: Method called, current is_paused = {self.is_paused}")
        self.is_paused = False
        print(f"🔧 RESUME DEBUG: After setting, is_paused = {self.is_paused}")
        print("▶️ Main assistant resumed from brainstorming session")
        
        # Check if audio stream is still healthy
        self.check_and_restart_audio_if_needed()
        
        # Use tkinter's thread-safe method for GUI updates
        def update_gui():
            try:
                self.text_display.insert(tk.END, "▶️ Main assistant resumed. You can continue with voice commands.\n")
                self.text_display.see(tk.END)
                print("✅ GUI updated - assistant resume message displayed")
            except Exception as e:
                print(f"⚠️ Error updating GUI in resume: {e}")
        
        # Schedule GUI update on main thread
        try:
            # Try to get the root window and schedule the update
            import tkinter as tk
            root = tk._default_root
            if root:
                root.after(0, update_gui)
            else:
                # Fallback to direct call if no root window
                update_gui()
        except Exception as e:
            print(f"⚠️ Error scheduling GUI update: {e}")
            # Final fallback
            update_gui()
    
    def check_and_restart_audio_if_needed(self):
        """Check if audio stream is healthy and restart if needed"""
        try:
            print("🔧 AUDIO CHECK: Verifying audio stream health...")
            
            # Check if audio stream is still active
            if not self.audio_stream or not self.audio_stream.is_active():
                print("⚠️ AUDIO CHECK: Audio stream is not active, attempting restart...")
                self.restart_audio_stream()
            elif not self.dg_connection:
                print("⚠️ AUDIO CHECK: Deepgram connection lost, attempting restart...")
                self.restart_audio_stream()
            else:
                print("✅ AUDIO CHECK: Audio stream appears healthy")
                
        except Exception as e:
            print(f"⚠️ AUDIO CHECK: Error checking audio health: {e}")
            print("🔄 AUDIO CHECK: Attempting restart due to error...")
            self.restart_audio_stream()
    
    def restart_audio_stream(self):
        """Restart the audio stream after brainstorming session"""
        try:
            print("🔄 AUDIO RESTART: Restarting audio stream...")
            
            # Stop current stream if it exists
            if self.audio_stream:
                try:
                    self.audio_stream.stop_stream()
                    self.audio_stream.close()
                    print("🛑 AUDIO RESTART: Stopped old audio stream")
                except:
                    pass
            
            # Finish current Deepgram connection if it exists
            if self.dg_connection:
                try:
                    self.dg_connection.finish()
                    print("🛑 AUDIO RESTART: Finished old Deepgram connection")
                except:
                    pass
            
            # Small delay to let audio device settle
            import time
            time.sleep(1.0)
            
            # Restart the audio stream in a new thread
            def restart_async():
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    loop.run_until_complete(self.start_audio_stream())
                finally:
                    loop.close()
            
            import threading
            restart_thread = threading.Thread(target=restart_async, daemon=True)
            restart_thread.start()
            
            print("✅ AUDIO RESTART: Audio stream restart initiated")
            
        except Exception as e:
            print(f"❌ AUDIO RESTART: Error restarting audio stream: {e}")
            import traceback
            print(f"🔍 AUDIO RESTART: Full traceback: {traceback.format_exc()}")
    
    def handle_free_time_week(self):
        """Handle weekly free time query"""
        if not self.check_authentication():
            return
            
        free_slots = get_free_slots_week()
        if free_slots:
            free_slots_by_day = {}
            self.text_display.insert(tk.END, "You are free at the following times this week:\n")
            print("You are free at the following times this week:")
            
            for start, end in free_slots:
                day = start.strftime("%A, %B %d")
                time_range = f"{start.strftime('%I:%M %p')} - {end.strftime('%I:%M %p')}"
                if day not in free_slots_by_day:
                    free_slots_by_day[day] = []
                free_slots_by_day[day].append(time_range)
                        
            for day, slots in free_slots_by_day.items():
                formatted_slots = ", ".join(slots)
                self.text_display.insert(tk.END, f"{day}: {formatted_slots}\n")
                print(f"{day}: {formatted_slots}")
        else:
            self.text_display.insert(tk.END, "You have no free time this week\n")
            print("You have no free time this week.")
                    
    def handle_free_time_today(self):
        """Handle today's free time query"""
        if not self.check_authentication():
            return
            
        free_slots = get_free_slots_today()
        if free_slots:
            self.text_display.insert(tk.END, "You are free at the following times today\n")
            print("You are free at the following times today:")
            for start, end in free_slots:
                self.text_display.insert(tk.END, f"{format_time(start)} - {format_time(end)}\n")
                print(f"{format_time(start)} - {format_time(end)}")
                        
    def handle_todays_events(self):
        """Handle today's events query"""
        if not self.check_authentication():
            return
            
        events_today = get_todays_events()
        self.text_display.insert(tk.END, "✅ Today's Events:\n")
        print("✅ Today's Events:")
        for event in events_today:
            event_summary = event.get("summary", "No Title")
            start_time = event.get("start_time", "Unknown Time")
            self.text_display.insert(tk.END, f"- {event_summary} at {start_time}\n")
            print(event)

    def handle_weeks_events(self):
        """Handle this week's events query"""
        if not self.check_authentication():
            return
            
        events_this_week = get_weeks_events()
        self.text_display.insert(tk.END, "✅ Events This Week:\n")
        print("✅ Events This Week:")
        for event in events_this_week:
            if event is None:
                continue
            event_summary = event.get("summary", "No Title")
            start_time = event.get("start_time", "Unknown Time")
            self.text_display.insert(tk.END, f"- {event_summary} at {start_time}\n")
            print(event)

    def handle_upcoming_events(self):
        """Handle upcoming events query"""
        if not self.check_authentication():
            return
            
        upcoming_events = get_events()
        self.text_display.insert(tk.END, "✅ Upcoming Events:\n")
        print("✅ Upcoming Events:")
        for event in upcoming_events: 
            event_summary = event.get("summary", "No Title")
            start_time = event.get("start_time", "Unknown Time")
            self.text_display.insert(tk.END, f"- {event_summary} at {start_time}\n")
            print(event)
                    
    def handle_question_response(self, question_text):
        """Handle the AI question response"""
        print(f"You asked: {question_text}")
        response = gen(question_text)
        self.text_display.insert(tk.END, f"AI Answer: {response}\n")
        print(f"AI Answer: {response}")
        self.waiting_for_question = False
    
    def handle_create_event(self, speech_text):
        """Handle event creation from speech"""
        if not self.check_authentication():
            return
        
        self.text_display.insert(tk.END, "🗓️ Creating calendar event...\n")
        self.text_display.see(tk.END)
        print("🗓️ Creating calendar event...")
        
        # Extract the event description (remove the trigger phrase)
        event_description = speech_text.lower().replace("create event", "").replace("schedule", "").strip()
        
        # Create event using AI parsing
        result = create_event_from_speech(event_description)
        
        if result:
            # Format the start time properly
            formatted_time = format_event_time(result['start_time'])
            
            self.text_display.insert(tk.END, f"✅ Event created: {result['summary']}\n")
            self.text_display.insert(tk.END, f"🕒 Time: {formatted_time}\n")
            if result.get('html_link'):
                self.text_display.insert(tk.END, f"🔗 Link: {result['html_link']}\n")
            self.text_display.see(tk.END)
            print(f"✅ Event created successfully")
        else:
            self.text_display.insert(tk.END, "❌ Failed to create event. Please try again.\n")
            self.text_display.see(tk.END)
            print("❌ Failed to create event")

    def handle_find_event(self, search_term):
        """Handle finding events by title"""
        if not self.check_authentication():
            return
        
        self.text_display.insert(tk.END, f"🔍 Searching for events: {search_term}\n")
        self.text_display.see(tk.END)
        print(f"🔍 Searching for events: {search_term}")
        
        events = find_events_by_title(search_term)
        
        if events:
            self.text_display.insert(tk.END, f"📅 Found {len(events)} matching events:\n")
            for event in events:
                self.text_display.insert(tk.END, f"- {event['summary']} at {event['start_time']}\n")
            self.text_display.see(tk.END)
        else:
            self.text_display.insert(tk.END, "❌ No matching events found.\n")
            self.text_display.see(tk.END)

    def handle_delete_event(self, search_term):
        """Handle deleting events"""
        if not self.check_authentication():
            return
        
        # First find the event
        events = find_events_by_title(search_term)
        
        if not events:
            self.text_display.insert(tk.END, "❌ No matching events found to delete.\n")
            self.text_display.see(tk.END)
            return
        
        if len(events) == 1:
            # Delete the single matching event
            event_id = events[0]['id']
            if delete_calendar_event(event_id):
                self.text_display.insert(tk.END, f"✅ Deleted event: {events[0]['summary']}\n")
                self.text_display.see(tk.END)
            else:
                self.text_display.insert(tk.END, "❌ Failed to delete event.\n")
                self.text_display.see(tk.END)
        else:
            # Multiple events found - show options
            self.text_display.insert(tk.END, f"🔍 Found {len(events)} matching events:\n")
            for i, event in enumerate(events):
                self.text_display.insert(tk.END, f"{i+1}. {event['summary']} at {event['start_time']}\n")
            self.text_display.insert(tk.END, "Please be more specific about which event to delete.\n")
            self.text_display.see(tk.END)
    
    def handle_move_event(self, speech_text):
        """Handle moving/rescheduling events"""
        if not self.check_authentication():
            return
        
        self.text_display.insert(tk.END, "📅 Processing event move request...\n")
        self.text_display.see(tk.END)
        print("📅 Processing event move request...")
        
        # Use AI to parse the move request
        result = parse_move_request(speech_text)
        
        if result:
            # Find the event to move
            events = find_events_by_title(result['event_search'])
            
            if not events:
                self.text_display.insert(tk.END, f"❌ No events found matching '{result['event_search']}'\n")
                self.text_display.see(tk.END)
                return
            
            if len(events) == 1:
                # Move the single matching event
                event = events[0]
                success = move_event_to_new_time(event, result['new_time'])
                if success:
                    self.text_display.insert(tk.END, f"✅ Moved '{event['summary']}' to {result['new_time']}\n")
                    self.text_display.see(tk.END)
                else:
                    self.text_display.insert(tk.END, "❌ Failed to move event.\n")
                    self.text_display.see(tk.END)
            else:
                # Multiple events found - show options
                self.text_display.insert(tk.END, f"🔍 Found {len(events)} matching events:\n")
                for i, event in enumerate(events):
                    self.text_display.insert(tk.END, f"{i+1}. {event['summary']} at {event['start_time']}\n")
                self.text_display.insert(tk.END, "Please be more specific about which event to move.\n")
                self.text_display.see(tk.END)
        else:
            self.text_display.insert(tk.END, "❌ Could not understand move request. Try: 'Move meeting to tomorrow 2 PM'\n")
            self.text_display.see(tk.END)

    async def start_audio_stream(self):
        """Start the audio stream and Deepgram transcription"""
        try:
            print("🎯 Starting audio stream setup...")
            
            if not self.deepgram:
                print("❌ Deepgram client not initialized")
                return
                
            print("🎤 Initializing PyAudio...")
            # Initialize PyAudio
            p = pyaudio.PyAudio()
            
            # Set up audio stream
            self.audio_stream = p.open(
                format=pyaudio.paInt16,
                channels=1,
                rate=16000,
                input=True,
                frames_per_buffer=1024
            )
            print("✅ Audio stream initialized")
            
            # Configure Deepgram options
            print("⚙️ Configuring Deepgram options...")
            options = LiveOptions(
                model="nova-3",  
                language="en-US",
                smart_format=True,
                interim_results=True,          # Changed to True like working code
                utterance_end_ms=1000,
                vad_events=True,
                endpointing=300,              # Added like working code
                punctuate=True,               # Added like working code
                diarize=True,                 # ← ENABLED for speaker diarization
                encoding="linear16",
                sample_rate=16000
            )
            print(f"✅ Options configured: model=nova-3, language=en-US, diarize=True")
            
            # Create a websocket connection using v4.1.0 pattern
            print("🌐 Creating WebSocket connection...")
            self.dg_connection = self.deepgram.listen.websocket.v("1")
            print(f"✅ WebSocket connection object created: {type(self.dg_connection)}")
            
            # Store reference to access VoiceAssistant instance from event handlers
            voice_assistant = self
            
            # Define event handlers using correct SDK 4.1.0 pattern
            print("📡 Setting up event handlers...")
            
            def on_open(self, open, **kwargs):
                print(f"🟢 Deepgram connection opened successfully!")

            def on_message(self, result, **kwargs):
                try:
                    sentence = result.channel.alternatives[0].transcript
                    if len(sentence) == 0:
                        print("📝 Empty transcript received")
                        return
                    
                    if result.is_final:
                        print(f"📝 Final transcript: '{sentence}'")
                        voice_assistant.process_transcript(result)
                    else:
                        print(f"📝 Interim transcript: '{sentence}'")
                except Exception as e:
                    print(f"❌ Error in transcript handler: {e}")

            def on_error(self, error, **kwargs):
                print(f"🔴 Deepgram error: {error}")

            def on_close(self, close, **kwargs):
                print(f"🔌 Deepgram connection closed")
            
            # Register event handlers using traditional callback pattern
            self.dg_connection.on(LiveTranscriptionEvents.Open, on_open)
            self.dg_connection.on(LiveTranscriptionEvents.Transcript, on_message)
            self.dg_connection.on(LiveTranscriptionEvents.Error, on_error)
            self.dg_connection.on(LiveTranscriptionEvents.Close, on_close)
            
            print("✅ Event handlers registered")
            
            # Start Deepgram connection with options
            print("🚀 Starting Deepgram connection...")
            start_result = self.dg_connection.start(options)
            print(f"🔍 Connection start result: {start_result}")
            
            if not start_result:
                print("❌ Failed to start Deepgram connection")
                return
            
            print("✅ Deepgram connection started successfully!")
            
            # Give the connection a moment to fully establish
            await asyncio.sleep(0.5)
            
            print("🎧 Starting audio streaming loop...")
            self.is_running = True
            
            # Audio streaming loop
            loop_count = 0
            while self.is_running:
                try:
                    # Read audio data
                    audio_data = self.audio_stream.read(1024, exception_on_overflow=False)
                    
                    # Send to Deepgram
                    if self.dg_connection:
                        self.dg_connection.send(audio_data)
                        
                        # Debug every 100 loops (roughly every second)
                        loop_count += 1
                        if loop_count % 100 == 0:
                            print(f"🔄 Audio streaming... (loop {loop_count}, sent {len(audio_data)} bytes)")
                        
                    # Small delay to prevent overwhelming the API
                    await asyncio.sleep(0.01)
                    
                except Exception as e:
                    print(f"❌ Error in audio loop: {e}")
                    # Don't break - continue trying to stream audio
                    await asyncio.sleep(0.1)  # Brief pause before retrying
            
        except Exception as e:
            print(f"❌ Error starting audio stream: {e}")
            import traceback
            print(f"🔍 Full traceback: {traceback.format_exc()}")
        finally:
            print("🧹 Cleaning up resources...")
            await self.cleanup()
    
    async def cleanup(self):
        """Clean up resources"""
        print("🛑 Stopping audio streaming...")
        self.is_running = False
        
        if self.dg_connection:
            print("🔌 Finishing Deepgram connection...")
            try:
                self.dg_connection.finish()
                print("✅ Deepgram connection finished")
            except Exception as e:
                print(f"❌ Error finishing Deepgram connection: {e}")
            
        if self.audio_stream:
            print("🎤 Closing audio stream...")
            try:
                self.audio_stream.stop_stream()
                self.audio_stream.close()
                print("✅ Audio stream closed")
            except Exception as e:
                print(f"❌ Error closing audio stream: {e}")
        
        print("✅ Cleanup completed")
    
    def stop_assistant(self):
        """Stop the assistant"""
        self.text_display.insert(tk.END, "Termination keyword detected. Stopping...\n")
        print("Termination keyword detected. Stopping...")
        self.is_running = False

def show_directions():
    """Creates a modern, editable directions window with custom voice commands."""
    
    # Load existing custom commands
    custom_commands = load_custom_commands()
    
    directions_window = tk.Toplevel(root)
    directions_window.title("Voice Commands & Settings")
    directions_window.geometry("400x600")
    directions_window.configure(bg="#0d1117")
    directions_window.resizable(True, True)
    
    # Center the window
    directions_window.geometry("+%d+%d" % (root.winfo_rootx() + 100, root.winfo_rooty() + 50))
    
    # Header Frame
    header_frame = tk.Frame(directions_window, bg="#161b22", height=70)
    header_frame.pack(fill=tk.X, padx=5, pady=5)
    header_frame.pack_propagate(False)
    
    header_inner = tk.Frame(header_frame, bg="#161b22")
    header_inner.pack(expand=True, fill=tk.BOTH)
    
    title = tk.Label(header_inner, text="🎤 Voice Commands & Settings", 
                    font=("Segoe UI", 18, "bold"), 
                    bg="#161b22", fg="#f0f6fc",
                    pady=15)
    title.pack()
    
    subtitle = tk.Label(header_inner, text="Customize your voice commands and learn how to use the assistant", 
                       font=("Segoe UI", 10), 
                       bg="#161b22", fg="#8b949e")
    subtitle.pack()
    
    # Create notebook for tabs
    notebook_frame = tk.Frame(directions_window, bg="#0d1117")
    notebook_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    # Tab buttons frame
    tab_frame = tk.Frame(notebook_frame, bg="#21262d", height=50)
    tab_frame.pack(fill=tk.X, pady=(0, 5))
    tab_frame.pack_propagate(False)
    
    # Content frame
    content_frame = tk.Frame(notebook_frame, bg="#21262d")
    content_frame.pack(fill=tk.BOTH, expand=True)
    
    # Tab content frames
    commands_frame = tk.Frame(content_frame, bg="#21262d")
    custom_frame = tk.Frame(content_frame, bg="#21262d")
    about_frame = tk.Frame(content_frame, bg="#21262d")
    
    current_tab = {"tab": "about"}
    
    def show_tab(tab_name, frame):
        # Hide all frames
        for frame_obj in [commands_frame, custom_frame, about_frame]:
            frame_obj.pack_forget()
        
        # Show selected frame
        frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        current_tab["tab"] = tab_name
        
        # Update tab button colors
        about_tab_btn.config(bg="#1f6feb" if tab_name == "about" else "#6f42c1")
        commands_tab_btn.config(bg="#1f6feb" if tab_name == "commands" else "#6f42c1")
        custom_tab_btn.config(bg="#1f6feb" if tab_name == "custom" else "#6f42c1")
    
    # Tab buttons
    about_tab_btn = tk.Button(tab_frame, text="📖 About", 
                              command=lambda: show_tab("about", about_frame),
                              font=("Segoe UI", 9, "bold"), 
                              bg="#1f6feb", fg="black",
                              relief=tk.FLAT, bd=0,
                              padx=15, pady=10,
                              cursor="hand2")
    about_tab_btn.pack(side=tk.LEFT, padx=5, pady=10)
    
    commands_tab_btn = tk.Button(tab_frame, text="📖 Shortcuts", 
                                command=lambda: show_tab("commands", commands_frame),
                                font=("Segoe UI", 9, "bold"), 
                                bg="#6f42c1", fg="black",
                                relief=tk.FLAT, bd=0,
                                padx=15, pady=10,
                                cursor="hand2")
    commands_tab_btn.pack(side=tk.LEFT, padx=5, pady=10)
    
    custom_tab_btn = tk.Button(tab_frame, text="⚙️ Custom", 
                              command=lambda: show_tab("custom", custom_frame),
                              font=("Segoe UI", 9, "bold"), 
                              bg="#6f42c1", fg="black",
                              relief=tk.FLAT, bd=0,
                              padx=15, pady=10,
                              cursor="hand2")
    custom_tab_btn.pack(side=tk.LEFT, padx=5, pady=10)
    
    # === COMMANDS TAB CONTENT ===
    # Scrollable frame for commands
    commands_canvas = tk.Canvas(commands_frame, bg="#21262d", highlightthickness=0)
    commands_scrollbar = tk.Scrollbar(commands_frame, orient="vertical", command=commands_canvas.yview,
                                     bg="#21262d", troughcolor="#0d1117", activebackground="#6f42c1")
    commands_scrollable = tk.Frame(commands_canvas, bg="#21262d")
    
    commands_scrollable.bind(
        "<Configure>",
        lambda e: commands_canvas.configure(scrollregion=commands_canvas.bbox("all"))
    )
    
    commands_canvas.create_window((0, 0), window=commands_scrollable, anchor="nw")
    commands_canvas.configure(yscrollcommand=commands_scrollbar.set)
    
    # Default commands content
    command_sections = [
        ("📅 Calendar Queries", [
            "When am I free this week?",
            "When am I free today?", 
            "What events do I have today?",
            "What events do I have this week?",
            "What events do I have coming up?"
        ]),
        ("📝 Event Management", [
            "Create event [description]",
            "Schedule [description]",
            "Example: 'Create event meeting with John tomorrow at 2 PM'"
        ]),
        ("🔍 Event Search", [
            "Find event [keyword]",
            "Search event [keyword]",
            "Example: 'Find event dentist'"
        ]),
        ("❌ Event Deletion", [
            "Delete event [keyword]",
            "Cancel event [keyword]",
            "Example: 'Delete event meeting'"
        ]),
        ("📅 Event Rescheduling", [
            "Move event [keyword] to [new time]",
            "Reschedule event [keyword] to [new time]",
            "Example: 'Move meeting to tomorrow 2 PM'"
        ]),
        ("🤖 AI Assistant", [
            "Question",
            "Terminate"
        ]),
        ("🧠 Brainstorming Sessions", [
            "Brainstorm [topic]",
            "Say 'summary' during session for recap",
            "Summaries appear in main GUI transcript",
            "Example: 'Brainstorm creative marketing ideas'",
            "Example: 'Brainstorm' (for open brainstorming)",
            "Note: Say 'end brainstorm' to end session"
        ]),
        ("🎯 Speaker Features", [
            "Speaker diarization automatically locks to first speaker",
            "Use 'Reset Speaker' button to re-identify primary speaker"
        ])
    ]
    
    for section_title, commands in command_sections:
        # Section header
        section_frame = tk.Frame(commands_scrollable, bg="#0d1117", relief=tk.FLAT, bd=1)
        section_frame.pack(fill=tk.X, pady=5)
        
        section_header = tk.Frame(section_frame, bg="#0d1117")
        section_header.pack(fill=tk.X, padx=15, pady=10)
        
        tk.Label(section_header, text=section_title, 
                font=("Segoe UI", 12, "bold"), 
                bg="#0d1117", fg="#f0f6fc").pack(anchor=tk.W)
        
        # Commands list
        for command in commands:
            cmd_frame = tk.Frame(section_frame, bg="#0d1117")
            cmd_frame.pack(fill=tk.X, padx=25, pady=2)
            
            tk.Label(cmd_frame, text=f"• {command}", 
                    font=("Consolas", 10), 
                    bg="#0d1117", fg="#8b949e",
                    anchor=tk.W).pack(anchor=tk.W)
    
    # === CUSTOM COMMANDS TAB CONTENT ===
    # Custom commands header
    custom_header = tk.Frame(custom_frame, bg="#21262d")
    custom_header.pack(fill=tk.X, pady=(0, 15))
    
    tk.Label(custom_header, text="⚙️ Customize Voice Shortcuts", 
            font=("Segoe UI", 14, "bold"), 
            bg="#21262d", fg="#f0f6fc").pack()
    
    tk.Label(custom_header, text="Add your own phrases that will trigger existing commands", 
            font=("Segoe UI", 10), 
            bg="#21262d", fg="#8b949e").pack(pady=(5, 0))
    
    # Custom commands scroll frame
    custom_canvas = tk.Canvas(custom_frame, bg="#21262d", highlightthickness=0)
    custom_scrollbar = tk.Scrollbar(custom_frame, orient="vertical", command=custom_canvas.yview,
                                   bg="#21262d", troughcolor="#0d1117", activebackground="#6f42c1")
    custom_scrollable = tk.Frame(custom_canvas, bg="#21262d")
    
    custom_scrollable.bind(
        "<Configure>",
        lambda e: custom_canvas.configure(scrollregion=custom_canvas.bbox("all"))
    )
    
    custom_canvas.create_window((0, 0), window=custom_scrollable, anchor="nw")
    custom_canvas.configure(yscrollcommand=custom_scrollbar.set)
    
    # Custom command categories with Entry widget tracking
    custom_categories = [
        ("📅 Calendar Events Today", "what events do i have today", [
            "What's on my schedule today?",
            "Today's agenda",
            "Show me today's calendar"
        ]),
        ("🆓 Free Time Today", "when am i free today", [
            "When can I schedule something today?",
            "What time slots are available?",
            "Show my free time"
        ]),
        ("📊 Weekly Events", "what events do i have this week", [
            "What's my week looking like?",
            "Weekly schedule",
            "Show this week's events"
        ]),
        ("📝 Create Event", "create event", [
            "Schedule a meeting",
            "Add to calendar",
            "Book an appointment"
        ]),
        ("🔍 Find Event", "find event", [
            "Look for event",
            "Search my calendar",
            "Where is my meeting"
        ]),
        ("❓ Ask Question", "question", [
            "I have a question",
            "Ask AI",
            "Help me with"
        ]),
        ("🧠 Start Brainstorming", "brainstorm", [
            "Let's brainstorm",
            "Creative session",
            "Generate ideas"
        ])
    ]
    
    # Dictionary to store Entry widgets for each command category
    entry_widgets = {}
    
    def add_phrase_to_category(base_command, phrases_frame):
        """Add a new phrase entry to a category"""
        phrase_frame = tk.Frame(phrases_frame, bg="#0d1117")
        phrase_frame.pack(fill=tk.X, pady=2)
        
        # Bullet point
        tk.Label(phrase_frame, text="•", 
                font=("Segoe UI", 10), 
                bg="#0d1117", fg="#8b949e").pack(side=tk.LEFT, padx=(0, 5))
        
        # Editable entry
        entry = tk.Entry(phrase_frame, 
                       font=("Consolas", 10), 
                       bg="#161b22", fg="#f0f6fc",
                       relief=tk.FLAT, bd=0,
                       insertbackground="#f0f6fc")
        entry.pack(fill=tk.X, padx=(0, 10))
        
        # Add to tracking
        if base_command not in entry_widgets:
            entry_widgets[base_command] = []
        entry_widgets[base_command].append(entry)
        
        return entry
    
    def save_custom_commands():
        """Save custom commands to user profile and reload them"""
        custom_commands = user_profile.load_custom_commands()
        
        # Collect data from all Entry widgets
        new_custom_commands = {}
        for base_command, entries in entry_widgets.items():
            phrases = []
            for entry in entries:
                phrase = entry.get().strip()
                if phrase:  # Only add non-empty phrases
                    phrases.append(phrase)
            new_custom_commands[base_command] = phrases
        
        try:
            # Save to user profile
            user_profile.save_custom_commands(new_custom_commands)
            
            # Show success message
            success_label = tk.Label(custom_scrollable, text="✅ Custom commands saved and reloaded!", 
                                   font=("Segoe UI", 10, "bold"), 
                                   bg="#21262d", fg="#3fb950")
            success_label.pack(pady=10)
            directions_window.after(3000, success_label.destroy)
            
            # Also show in main transcript
            text_display.insert(tk.END, f"✅ Custom commands updated! Saved {len([p for phrases in new_custom_commands.values() for p in phrases])} total phrases.\n")
            text_display.see(tk.END)
            
            print(f"✅ Custom commands saved: {new_custom_commands}")
            
        except Exception as e:
            error_label = tk.Label(custom_scrollable, text=f"❌ Error saving: {e}", 
                                 font=("Segoe UI", 10, "bold"), 
                                 bg="#21262d", fg="#f85149")
            error_label.pack(pady=10)
            directions_window.after(5000, error_label.destroy)
            print(f"❌ Error saving custom commands: {e}")
    
    # Create editable sections for each category
    for category, base_command, default_phrases in custom_categories:
        # Initialize entry tracking for this command
        entry_widgets[base_command] = []
        
        # Category card
        category_card = tk.Frame(custom_scrollable, bg="#0d1117", relief=tk.FLAT, bd=1)
        category_card.pack(fill=tk.X, pady=8)
        
        # Category header
        cat_header = tk.Frame(category_card, bg="#0d1117")
        cat_header.pack(fill=tk.X, padx=15, pady=10)
        
        tk.Label(cat_header, text=category, 
                font=("Segoe UI", 11, "bold"), 
                bg="#0d1117", fg="#f0f6fc").pack(anchor=tk.W)
        
        tk.Label(cat_header, text=f"Triggers: '{base_command}' command", 
                font=("Segoe UI", 9), 
                bg="#0d1117", fg="#8b949e").pack(anchor=tk.W, pady=(2, 0))
        
        # Editable phrases
        phrases_frame = tk.Frame(category_card, bg="#0d1117")
        phrases_frame.pack(fill=tk.X, padx=25, pady=(0, 15))
        
        # Add default phrases or load existing ones
        phrases_to_load = default_phrases
        if base_command in custom_commands and custom_commands[base_command]:
            phrases_to_load = custom_commands[base_command]
        
        for phrase in phrases_to_load:
            entry = add_phrase_to_category(base_command, phrases_frame)
            entry.insert(0, phrase)
        
        # Add 2 empty slots for new phrases
        for _ in range(2):
            add_phrase_to_category(base_command, phrases_frame)
        
        # Add new phrase button
        add_btn = tk.Button(phrases_frame, text="+ Add Phrase", 
                           command=lambda cmd=base_command, frame=phrases_frame: add_phrase_to_category(cmd, frame),
                           font=("Segoe UI", 9), 
                           bg="#238636", fg="black",
                           relief=tk.FLAT, bd=0,
                           padx=10, pady=4,
                           cursor="hand2")
        add_btn.pack(anchor=tk.W, pady=(5, 0))
    
    # Save button
    save_frame = tk.Frame(custom_scrollable, bg="#21262d")
    save_frame.pack(fill=tk.X, pady=20)
    
    save_btn = tk.Button(save_frame, text="💾 Save Custom Commands", 
                        command=save_custom_commands,
                        font=("Segoe UI", 12, "bold"), 
                        bg="#1f6feb", fg="black",
                        relief=tk.FLAT, bd=0,
                        padx=30, pady=12,
                        cursor="hand2")
    save_btn.pack()
    
    # === ABOUT TAB CONTENT ===
    # About tab header
    about_header = tk.Frame(about_frame, bg="#21262d")
    about_header.pack(fill=tk.X, pady=(0, 15))
    
    tk.Label(about_header, text="🤖 AI Voice Assistant", 
            font=("Segoe UI", 16, "bold"), 
            bg="#21262d", fg="#f0f6fc").pack()
    
    tk.Label(about_header, text="Your intelligent voice-powered productivity companion", 
            font=("Segoe UI", 11), 
            bg="#21262d", fg="#8b949e").pack(pady=(5, 0))
    
    # About content scroll frame
    about_canvas = tk.Canvas(about_frame, bg="#21262d", highlightthickness=0)
    about_scrollbar = tk.Scrollbar(about_frame, orient="vertical", command=about_canvas.yview,
                                  bg="#21262d", troughcolor="#0d1117", activebackground="#6f42c1")
    about_scrollable = tk.Frame(about_canvas, bg="#21262d")
    
    about_scrollable.bind(
        "<Configure>",
        lambda e: about_canvas.configure(scrollregion=about_canvas.bbox("all"))
    )
    
    about_canvas.create_window((0, 0), window=about_scrollable, anchor="nw")
    about_canvas.configure(yscrollcommand=about_scrollbar.set)
    
    # About sections
    about_sections = [
        ("🎯 Welcome to Your AI Voice Assistant", [
            "This application was built to transform the way you interact with technology through natural speech.",
            "We designed this tool specifically for busy professionals, creatives, and anyone who thinks faster than they type.",
            "Your voice becomes a powerful interface for managing calendars, brainstorming ideas, and accessing AI assistance.",
            "This bridges the gap between natural human communication and digital productivity tools."
        ]),
        
        ("⚡ Core Functionality", [
            "📅 Smart Calendar Management - Create, search, move, and delete events by voice",
            "🧠 AI-Powered Brainstorming - Voice conversations with creative AI partner",
            "🎤 Speaker Diarization - Focuses on your voice in multi-speaker environments",
            "🤖 Intelligent Q&A - Ask questions and get AI-powered answers instantly",
            "⚙️ Custom Voice Commands - Personalize phrases to trigger any function",
            "📝 Session Summaries - Automatic brainstorming session documentation",
            "🔗 Google Calendar Integration - Full two-way sync with your existing calendar"
        ]),
        
        ("🛠️ Technology Stack", [
            "🎙️ Deepgram Nova-3 - Advanced speech-to-text with speaker diarization",
            "🧠 OpenAI GPT-4 - Powering intelligent responses and brainstorming",
            "📅 Google Calendar API - Seamless calendar integration and management",
            "🐍 Python + Tkinter - Cross-platform desktop application framework",
            "🎵 PyAudio - Real-time audio capture and processing",
            "⚡ Asyncio - Efficient asynchronous audio streaming",
            "🔒 OAuth 2.0 - Secure Google account authentication"
        ]),
        
        ("🚀 Getting Started", [
            "1. Connect Google Calendar using the 'Connect' button below",
            "2. Set your timezone for accurate event scheduling",
            "3. Click 'Start Listening' to begin voice interaction",
            "4. Try saying 'What events do I have today?' or 'Brainstorm marketing ideas'",
            "5. Customize voice commands in the 'Custom Commands' tab",
            "6. Use 'Help' for detailed command reference"
        ]),
        
        ("💡 Advanced Features", [
            "🎯 Smart Speaker Lock - Automatically focuses on the first speaker",
            "📊 Voice Analytics - Track brainstorming session statistics",
            "💾 Transcript Export - Save conversations to Word documents",
            "🔄 Session Management - Seamless switching between modes",
            "📝 Real-time Feedback - Live status updates and confirmations",
            "🎨 Modern UI - Dark theme optimized for extended use"
        ]),
        
        ("🔐 Privacy & Security", [
            "🔒 Local Processing - Your conversations stay on your device when possible",
            "🛡️ Secure Authentication - OAuth 2.0 for Google Calendar access",
            "🗑️ Data Control - Clear transcripts and revoke access anytime",
            "🔑 API Key Management - Your Deepgram and OpenAI keys are encrypted",
            "📋 No Data Mining - We don't store or analyze your personal conversations"
        ]),
        
        ("🎯 Perfect For", [
            "📋 Project Managers - Quick meeting scheduling and task creation",
            "🎨 Creative Professionals - Voice brainstorming and idea generation",
            "👥 Team Leaders - Efficient calendar coordination and planning",
            "🏢 Executives - Hands-free calendar management during busy days",
            "📚 Students - Voice-controlled study scheduling and research assistance",
            "🚀 Entrepreneurs - Rapid idea capture and schedule optimization"
        ]),
        
        ("🔮 Future Roadmap", [
            "📱 Mobile App - iOS and Android versions in development",
            "🔗 More Integrations - Slack, Teams, Notion, and Trello support",
            "🌐 Multi-language - Support for Spanish, French, German, and more",
            "🤝 Team Features - Shared brainstorming sessions and collaborative calendars",
            "📊 Analytics Dashboard - Productivity insights and usage statistics",
            "🎵 Voice Cloning - Personalized AI voice responses"
        ])
    ]
    
    for section_title, items in about_sections:
        # Section card
        section_card = tk.Frame(about_scrollable, bg="#0d1117", relief=tk.FLAT, bd=1)
        section_card.pack(fill=tk.X, pady=8)
        
        # Section header
        section_header = tk.Frame(section_card, bg="#0d1117")
        section_header.pack(fill=tk.X, padx=15, pady=12)
        
        tk.Label(section_header, text=section_title, 
                font=("Segoe UI", 12, "bold"), 
                bg="#0d1117", fg="#f0f6fc").pack(anchor=tk.W)
        
        # Items list
        items_frame = tk.Frame(section_card, bg="#0d1117")
        items_frame.pack(fill=tk.X, padx=25, pady=(0, 15))
        
        for item in items:
            item_frame = tk.Frame(items_frame, bg="#0d1117")
            item_frame.pack(fill=tk.X, pady=3)
            
            # Special handling for the welcome section (no bullet points)
            if section_title == "🎯 Welcome to Your AI Voice Assistant":
                tk.Label(item_frame, text=item, 
                        font=("Segoe UI", 10), 
                        bg="#0d1117", fg="#f0f6fc",
                        anchor=tk.W, wraplength=350, justify=tk.LEFT).pack(anchor=tk.W, pady=2)
            elif item.startswith(("1.", "2.", "3.", "4.", "5.", "6.")):
                # Numbered steps - use different styling
                tk.Label(item_frame, text=item, 
                        font=("Segoe UI", 10, "bold"), 
                        bg="#0d1117", fg="#3fb950",
                        anchor=tk.W, wraplength=350, justify=tk.LEFT).pack(anchor=tk.W)
            elif " - " in item:
                # Feature descriptions - split into title and description
                parts = item.split(" - ", 1)
                title = parts[0]
                desc = parts[1] if len(parts) > 1 else ""
                
                # Title in bold
                title_label = tk.Label(item_frame, text=f"• {title}", 
                                     font=("Segoe UI", 10, "bold"), 
                                     bg="#0d1117", fg="#58a6ff",
                                     anchor=tk.W)
                title_label.pack(anchor=tk.W)
                
                # Description in regular text
                if desc:
                    desc_label = tk.Label(item_frame, text=f"  {desc}", 
                                        font=("Segoe UI", 9), 
                                        bg="#0d1117", fg="#8b949e",
                                        anchor=tk.W, wraplength=330, justify=tk.LEFT)
                    desc_label.pack(anchor=tk.W, padx=(20, 0))
            else:
                # Regular bullet points
                tk.Label(item_frame, text=f"• {item}", 
                        font=("Segoe UI", 10), 
                        bg="#0d1117", fg="#8b949e",
                        anchor=tk.W, wraplength=350, justify=tk.LEFT).pack(anchor=tk.W)
    
    # Version and credits footer for about tab
    about_footer = tk.Frame(about_scrollable, bg="#0d1117")
    about_footer.pack(fill=tk.X, pady=20)
    
    tk.Label(about_footer, text="🚀 Version 1.0 - Built with ❤️ for productivity enthusiasts", 
            font=("Segoe UI", 10, "italic"), 
            bg="#0d1117", fg="#6f42c1").pack()
    
    tk.Label(about_footer, text="Powered by Deepgram, OpenAI, and Google Calendar APIs", 
            font=("Segoe UI", 8), 
            bg="#0d1117", fg="#8b949e").pack(pady=(5, 0))
    
    # Pack scrollable areas
    def pack_commands_tab():
        commands_canvas.pack(side="left", fill="both", expand=True)
        commands_scrollbar.pack(side="right", fill="y")
    
    def pack_custom_tab():
        custom_canvas.pack(side="left", fill="both", expand=True)
        custom_scrollbar.pack(side="right", fill="y")
    
    def pack_about_tab():
        about_canvas.pack(side="left", fill="both", expand=True)
        about_scrollbar.pack(side="right", fill="y")
    
    # Initially show about tab
    show_tab("about", about_frame)
    pack_about_tab()
    
    # Update tab switching to handle scrollbars
    def enhanced_show_tab(tab_name, frame):
        # Hide all frames and scrollbars
        for frame_obj in [commands_frame, custom_frame, about_frame]:
            frame_obj.pack_forget()
        commands_scrollbar.pack_forget()
        custom_scrollbar.pack_forget()
        commands_canvas.pack_forget()
        custom_canvas.pack_forget()
        about_scrollbar.pack_forget()
        about_canvas.pack_forget()
        
        # Show selected frame and scrollbar
        frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        if tab_name == "about":
            pack_about_tab()
        elif tab_name == "commands":
            pack_commands_tab()
        elif tab_name == "custom":
            pack_custom_tab()
        
        current_tab["tab"] = tab_name
        
        # Update tab button colors
        about_tab_btn.config(bg="#1f6feb" if tab_name == "about" else "#6f42c1")
        commands_tab_btn.config(bg="#1f6feb" if tab_name == "commands" else "#6f42c1")
        custom_tab_btn.config(bg="#1f6feb" if tab_name == "custom" else "#6f42c1")
    
    # Update button commands
    about_tab_btn.config(command=lambda: enhanced_show_tab("about", about_frame))
    commands_tab_btn.config(command=lambda: enhanced_show_tab("commands", commands_frame))
    custom_tab_btn.config(command=lambda: enhanced_show_tab("custom", custom_frame))
    
    # Footer with close button
    footer_frame = tk.Frame(directions_window, bg="#161b22", height=50)
    footer_frame.pack(fill=tk.X, padx=5, pady=5)
    footer_frame.pack_propagate(False)
    
    close_btn = tk.Button(footer_frame, text="✕ Close", 
                         command=directions_window.destroy,
                         font=("Segoe UI", 10, "bold"), 
                         bg="#da3633", fg="black",
                         relief=tk.FLAT, bd=0,
                         padx=20, pady=8,
                         cursor="hand2")
    close_btn.pack(side=tk.RIGHT, padx=15, pady=10)
    
    # Bind mouse wheel to canvas
    def on_mousewheel(event):
        if current_tab["tab"] == "about":
            about_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        elif current_tab["tab"] == "commands":
            commands_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        elif current_tab["tab"] == "custom":
            custom_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
    
    directions_window.bind("<MouseWheel>", on_mousewheel)

def authenticate_google():
    """Handle Google Calendar authentication"""
    if authenticate():
        update_auth_status()
        text_display.insert(tk.END, "✅ Google Calendar authentication successful!\n")
    else:
        text_display.insert(tk.END, "❌ Google Calendar authentication failed.\n")

def revoke_google_access():
    """Handle revoking Google Calendar access"""
    revoke_access()
    update_auth_status()
    text_display.insert(tk.END, "🔓 Google Calendar access revoked.\n")

def update_auth_status():
    """Update the authentication status display with modern styling"""
    if is_authenticated():
        auth_status_label.config(text="✅ Google Calendar: Connected", fg="#3fb950")
        auth_button.config(text="🔓 Disconnect", command=revoke_google_access, bg="#da3633")
    else:
        auth_status_label.config(text="❌ Google Calendar: Not Connected", fg="#f85149")
        auth_button.config(text="🔗 Connect Google Calendar", command=authenticate_google, bg="#238636")
    
def start_assistant():
    """Start the voice assistant with UI updates"""
    global current_assistant
    current_assistant = VoiceAssistant(text_display)
    
    # Register GUI callback for brainstorming summaries
    set_gui_callback(display_brainstorm_message)
    
    # Register resume callback for when brainstorming ends
    set_assistant_resume_callback(current_assistant.resume_assistant)
    print(f"🔧 CALLBACK DEBUG: Resume callback registered for assistant instance {id(current_assistant)}")
    print(f"🔧 CALLBACK DEBUG: Initial is_paused state = {current_assistant.is_paused}")
    
    # Update UI
    main_button.config(text="🛑 Stop Listening", command=stop_assistant_ui, bg="#da3633")
    update_voice_status("Initializing...", True)
    reset_speaker_button.config(state=tk.NORMAL)
    
    def run_async():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            loop.run_until_complete(current_assistant.start_audio_stream())
        finally:
            loop.close()
    
    thread = threading.Thread(target=run_async)
    thread.daemon = True
    thread.start()

    # Welcome message with better formatting
    text_display.insert(tk.END, "🤖 AI Voice Assistant - Ready!\n")
    text_display.insert(tk.END, "=" * 40 + "\n")
    text_display.insert(tk.END, "🎤 Speaker diarization enabled\n")
    text_display.insert(tk.END, "🗣️ Say 'Terminate' to stop\n")
    text_display.insert(tk.END, "🧠 Say 'Brainstorm' to start creative session\n")
    text_display.insert(tk.END, "📋 Summaries from brainstorming will appear here\n")
    text_display.insert(tk.END, "=" * 40 + "\n\n")
    text_display.see(tk.END)
    
    # Update status after a short delay
    root.after(2000, lambda: update_voice_status("Listening...", True))

def stop_assistant_ui():
    """Stop assistant with UI updates"""
    global current_assistant
    if current_assistant:
        current_assistant.stop_assistant()
        current_assistant = None
    
    # Clear callbacks to prevent memory leaks
    set_gui_callback(None)
    set_assistant_resume_callback(None)
    
    # Reset UI
    main_button.config(text="🎤 Start Listening", command=start_assistant, bg="#1f6feb")
    update_voice_status("Ready to Listen", False)
    update_speaker_lock_status(False)
    reset_speaker_button.config(state=tk.DISABLED)

def reset_speaker():
    """Reset the speaker lock with UI updates"""
    global current_assistant
    if current_assistant:
        current_assistant.reset_speaker_lock()
        update_speaker_lock_status(False)
        text_display.insert(tk.END, "🔄 Speaker lock reset\n")
        text_display.see(tk.END)
    else:
        text_display.insert(tk.END, "❌ Assistant not running. Start assistant first.\n")
        text_display.see(tk.END)

def get_downloads_folder():
    """Get the user's Downloads folder path, cross-platform compatible"""
    try:
        downloads_path = Path.home() / "Downloads"
        # Create Downloads folder if it doesn't exist
        downloads_path.mkdir(exist_ok=True)
        return downloads_path
    except Exception as e:
        print(f"❌ Error accessing Downloads folder: {e}")
        # Fallback to home directory
        return Path.home()

def save_transcript_to_word():
    """Save the current transcript to a Word document in Downloads folder"""
    try:
        # Try to import python-docx
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # Show error message in transcript
            text_display.insert(tk.END, "❌ python-docx not installed. Installing...\n")
            text_display.see(tk.END)
            
            # Try to install python-docx
            import subprocess
            import sys
            result = subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], 
                                  capture_output=True, text=True)
            
            if result.returncode == 0:
                text_display.insert(tk.END, "✅ python-docx installed successfully!\n")
                text_display.see(tk.END)
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            else:
                text_display.insert(tk.END, "❌ Failed to install python-docx. Please install manually.\n")
                text_display.see(tk.END)
                return
        
        # Get transcript content
        transcript_content = text_display.get(1.0, tk.END).strip()
        
        if not transcript_content:
            text_display.insert(tk.END, "❌ No transcript content to save.\n")
            text_display.see(tk.END)
            return
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('AI Voice Assistant Transcript', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        from datetime import datetime
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("=" * 50)
        doc.add_paragraph("")
        
        # Add transcript content
        # Split by lines and format nicely
        lines = transcript_content.split('\n')
        current_paragraph = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_paragraph("")  # Add blank line
            elif line.startswith("="):
                # Section separator
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                doc.add_paragraph(line)
            elif line.startswith("🤖") or line.startswith("✅") or line.startswith("❌") or line.startswith("🎤"):
                # System messages - make them bold
                if current_paragraph:
                    doc.add_paragraph(current_paragraph)
                    current_paragraph = ""
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            else:
                # Regular content
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
        
        # Add any remaining content
        if current_paragraph:
            doc.add_paragraph(current_paragraph)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"voice_assistant_transcript_{timestamp}.docx"
        
        # Get Downloads folder path
        downloads_folder = get_downloads_folder()
        file_path = downloads_folder / filename
        
        # Save document
        doc.save(str(file_path))
        
        # Show success message
        text_display.insert(tk.END, f"✅ Transcript saved as: {filename}\n")
        text_display.insert(tk.END, f"📁 Location: {downloads_folder}\n")
        text_display.see(tk.END)
        
        print(f"✅ Transcript saved to: {file_path}")
        
    except Exception as e:
        error_msg = f"❌ Error saving transcript: {str(e)}\n"
        text_display.insert(tk.END, error_msg)
        text_display.see(tk.END)
        print(f"Error saving transcript: {e}")

def on_timezone_change(event):
    """Handle timezone selection change"""
    selected_timezone = timezone_var.get()
    if selected_timezone and selected_timezone != "Select Timezone":
        success = set_user_timezone(selected_timezone)
        if success:
            text_display.insert(tk.END, f"🌍 Timezone set to: {selected_timezone}\n")
            text_display.see(tk.END)
        else:
            text_display.insert(tk.END, f"❌ Failed to set timezone: {selected_timezone}\n")
            text_display.see(tk.END)

# GUI Setup - Modern Compact Voice Assistant
root = tk.Tk()
root.title("AI Voice Assistant")
root.geometry("400x600")
root.configure(bg="#0d1117")  # GitHub dark theme
root.resizable(True, True)

# Make window always on top (optional)
root.attributes('-topmost', False)

# Custom styling
def create_gradient_frame(parent, color1, color2):
    """Create a gradient-like effect using frames"""
    frame = tk.Frame(parent, bg=color1)
    return frame

# Header Frame with modern styling
header_frame = tk.Frame(root, bg="#161b22", height=80)
header_frame.pack(fill=tk.X, padx=2, pady=2)
header_frame.pack_propagate(False)

# App title with modern styling
title_frame = tk.Frame(header_frame, bg="#161b22")
title_frame.pack(expand=True, fill=tk.BOTH)

app_title = tk.Label(title_frame, text="🎤 AI Voice Assistant", 
                    font=("Segoe UI", 16, "bold"), 
                    bg="#161b22", fg="#f0f6fc",
                    pady=10)
app_title.pack()

subtitle = tk.Label(title_frame, text="Powered by Deepgram", 
                   font=("Segoe UI", 9), 
                   bg="#161b22", fg="#8b949e")
subtitle.pack()

# Status indicator frame
status_frame = tk.Frame(root, bg="#0d1117", height=60)
status_frame.pack(fill=tk.X, padx=10, pady=5)
status_frame.pack_propagate(False)

# Voice status indicator
voice_status_frame = tk.Frame(status_frame, bg="#0d1117")
voice_status_frame.pack(side=tk.LEFT, fill=tk.Y)

voice_indicator = tk.Label(voice_status_frame, text="⚫", font=("Arial", 20), 
                          bg="#0d1117", fg="#f85149")  # Red when not listening
voice_indicator.pack(side=tk.LEFT, padx=5)

voice_status_label = tk.Label(voice_status_frame, text="Ready to Listen", 
                             font=("Segoe UI", 10, "bold"), 
                             bg="#0d1117", fg="#f0f6fc")
voice_status_label.pack(side=tk.LEFT, padx=5)

# Speaker lock indicator
speaker_frame = tk.Frame(status_frame, bg="#0d1117")
speaker_frame.pack(side=tk.RIGHT, fill=tk.Y)

speaker_lock_label = tk.Label(speaker_frame, text="🔓 No Speaker Lock", 
                             font=("Segoe UI", 9), 
                             bg="#0d1117", fg="#8b949e")
speaker_lock_label.pack(side=tk.RIGHT, padx=5)

# Authentication status with modern card design
auth_card = tk.Frame(root, bg="#21262d", relief=tk.FLAT, bd=1)
auth_card.pack(fill=tk.X, padx=10, pady=5)

auth_inner = tk.Frame(auth_card, bg="#21262d")
auth_inner.pack(fill=tk.X, padx=15, pady=10)

# First row: Auth status
auth_status_label = tk.Label(auth_inner, text="❌ Google Calendar: Not Connected", 
                           font=("Segoe UI", 10), bg="#21262d", fg="#f85149")
auth_status_label.pack(anchor=tk.W)

# Second row: Auth button and timezone dropdown
auth_controls_frame = tk.Frame(auth_inner, bg="#21262d")
auth_controls_frame.pack(fill=tk.X, pady=(5, 0))

auth_button = tk.Button(auth_controls_frame, text="🔗 Connect Google Calendar", 
                       command=authenticate_google,
                       font=("Segoe UI", 9, "bold"), 
                       bg="#238636", fg="black", 
                       relief=tk.FLAT, bd=0,
                       padx=15, pady=8,
                       cursor="hand2")
auth_button.pack(side=tk.LEFT)

# Timezone dropdown
timezone_frame = tk.Frame(auth_controls_frame, bg="#21262d")
timezone_frame.pack(side=tk.RIGHT, padx=(10, 0))

timezone_label = tk.Label(timezone_frame, text="🌍 Timezone:", 
                         font=("Segoe UI", 9), bg="#21262d", fg="#f0f6fc")
timezone_label.pack(side=tk.LEFT, padx=(0, 5))

# Common timezones list
common_timezones = [
    "Select Timezone",
    "US/Eastern",
    "US/Central", 
    "US/Mountain",
    "US/Pacific",
    "Europe/London",
    "Europe/Paris",
    "Europe/Berlin",
    "Asia/Tokyo",
    "Asia/Shanghai",
    "Australia/Sydney",
    "UTC"
]

timezone_var = tk.StringVar(value="Select Timezone")
timezone_dropdown = ttk.Combobox(timezone_frame, textvariable=timezone_var, 
                                values=common_timezones,
                                font=("Segoe UI", 8),
                                width=15,
                                state="readonly")
timezone_dropdown.pack(side=tk.LEFT)
timezone_dropdown.bind("<<ComboboxSelected>>", on_timezone_change)

# Control buttons with modern design
control_card = tk.Frame(root, bg="#21262d", relief=tk.FLAT, bd=1)
control_card.pack(fill=tk.X, padx=10, pady=5)

control_inner = tk.Frame(control_card, bg="#21262d")
control_inner.pack(fill=tk.X, padx=15, pady=15)

# Main control button (large, prominent)
main_button = tk.Button(control_inner, text="🎤 Start Listening", 
                       command=start_assistant,
                       font=("Segoe UI", 12, "bold"), 
                       bg="#1f6feb", fg="black",
                       relief=tk.FLAT, bd=0,
                       padx=20, pady=12,
                       cursor="hand2")
main_button.pack(fill=tk.X, pady=(0, 10))

# Secondary buttons row
secondary_frame = tk.Frame(control_inner, bg="#21262d")
secondary_frame.pack(fill=tk.X)

reset_speaker_button = tk.Button(secondary_frame, text="🔄 Reset Speaker", 
                                command=reset_speaker,
                                font=("Segoe UI", 9), 
                                bg="#6f42c1", fg="Black",
                                relief=tk.FLAT, bd=0,
                                padx=10, pady=6,
                                cursor="hand2")
reset_speaker_button.pack(side=tk.LEFT, padx=(0, 5))

directions_button = tk.Button(secondary_frame, text="📖 Help", 
                             command=show_directions,
                             font=("Segoe UI", 9), 
                             bg="#6f42c1", fg="black",
                             relief=tk.FLAT, bd=0,
                             padx=10, pady=6,
                             cursor="hand2")
directions_button.pack(side=tk.LEFT, padx=5)

# Save transcript feature
save_transcript_button = tk.Button(secondary_frame, text="💾 Save Transcript", 
                                command=lambda: save_transcript_to_word(),
                                font=("Segoe UI", 9), 
                                bg="#6f42c1", fg="black",
                                relief=tk.FLAT, bd=0,
                                padx=10, pady=6,
                                cursor="hand2")
save_transcript_button.pack(side=tk.RIGHT)

# Transcript display with modern styling
transcript_card = tk.Frame(root, bg="#21262d", relief=tk.FLAT, bd=1)
transcript_card.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

transcript_header = tk.Frame(transcript_card, bg="#21262d")
transcript_header.pack(fill=tk.X, padx=15, pady=(10, 0))

transcript_title = tk.Label(transcript_header, text="💬 Conversation", 
                           font=("Segoe UI", 11, "bold"), 
                           bg="#21262d", fg="#f0f6fc")
transcript_title.pack(side=tk.LEFT)

# Clear button for transcript
clear_button = tk.Button(transcript_header, text="🗑️ Clear", 
                        command=lambda: text_display.delete(1.0, tk.END),
                        font=("Segoe UI", 8), 
                        bg="#da3633", fg="black",
                        relief=tk.FLAT, bd=0,
                        padx=8, pady=4,
                        cursor="hand2")
clear_button.pack(side=tk.RIGHT)

# Text display with scrollbar
text_frame = tk.Frame(transcript_card, bg="#21262d")
text_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(5, 15))

text_display = tk.Text(text_frame, 
                      wrap=tk.WORD, 
                      font=("Consolas", 10), 
                      bg="#0d1117", fg="#f0f6fc",
                      insertbackground="#f0f6fc",
                      selectbackground="#1f6feb",
                      relief=tk.FLAT, bd=0,
                      padx=10, pady=10)
text_display.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Custom scrollbar
scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_display.yview,
                        bg="#21262d", troughcolor="#0d1117", 
                        activebackground="#6f42c1")
text_display.configure(yscrollcommand=scrollbar.set)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

# Footer with status and controls
footer_frame = tk.Frame(root, bg="#161b22", height=40)
footer_frame.pack(fill=tk.X, padx=2, pady=2)
footer_frame.pack_propagate(False)

# Mini controls in footer
footer_inner = tk.Frame(footer_frame, bg="#161b22")
footer_inner.pack(expand=True, fill=tk.BOTH)

# Connection status
connection_status = tk.Label(footer_inner, text="🔴 Disconnected", 
                           font=("Segoe UI", 8), 
                           bg="#161b22", fg="#8b949e")
connection_status.pack(side=tk.LEFT, padx=10, pady=10)

# Window controls
minimize_button = tk.Button(footer_inner, text="−", 
                           command=root.iconify,
                           font=("Arial", 12, "bold"), 
                           bg="#161b22", fg="#8b949e",
                           relief=tk.FLAT, bd=0,
                           width=3, height=1,
                           cursor="hand2")
minimize_button.pack(side=tk.RIGHT, padx=2, pady=8)

# Enhanced functions for UI updates
def update_voice_status(status, is_listening=False):
    """Update voice status indicator"""
    if is_listening:
        voice_indicator.config(fg="#3fb950", text="🟢")  # Green when listening
        voice_status_label.config(text=status, fg="#3fb950")
        connection_status.config(text="🟢 Connected", fg="#3fb950")
    else:
        voice_indicator.config(fg="#f85149", text="⚫")  # Red when not listening
        voice_status_label.config(text=status, fg="#f0f6fc")
        connection_status.config(text="🔴 Disconnected", fg="#f85149")

def update_speaker_lock_status(locked=False, speaker_id=None):
    """Update speaker lock indicator"""
    if locked and speaker_id is not None:
        speaker_lock_label.config(text=f"🔒 Speaker {speaker_id} Locked", fg="#3fb950")
    else:
        speaker_lock_label.config(text="🔓 No Speaker Lock", fg="#8b949e")

# Initialize UI
update_auth_status()

# Load custom commands on startup
load_custom_commands()

# Add some style enhancements
def on_enter(event):
    """Button hover effect"""
    event.widget.config(bg=event.widget.cget('bg').replace('#', '#').replace('1f6feb', '0969da'))

def on_leave(event):
    """Button hover effect"""
    event.widget.config(bg=event.widget.cget('bg').replace('#', '#').replace('0969da', '1f6feb'))

# Bind hover effects to main buttons
for button in [main_button, auth_button, reset_speaker_button, directions_button, save_transcript_button]:
    button.bind("<Enter>", on_enter)
    button.bind("<Leave>", on_leave)

root.mainloop()
